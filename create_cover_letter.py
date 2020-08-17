
import click
from docx import Document



def if_docx(job, company):
    hi = "Hi, "
    first_paragraph = f"My name is Kim Kiamco,  I am excited for the opportunity to apply for the {job} position. I will be an asset to the team because of my commitment to building quality web applications. I am a part-time student at coding bootcamp (Lambda school) taking the Full-Stack Web Development track and will be finishing the program soon. Currently, I work as a part time Operations engineer at Menlo Security, a cyber security company. I have been looking for a full time position and thought {company} would be a great company to work for. My skills from Lambda school, working experience from Menlo Security and my adaptability in any environment will allow me to excel at {company}. My main tech stack when working on web applications is Javascript, React, Node.js, and PostgreSQL but I am willing and excited to learn new languages if needed. I am familiar with DevOps tools and practices like Docker, AWS and can write python scripts that can help with deployment automation."
    second_paragraph = f"I have attached my resume and I believe that my experience and education would make me an asset at {company},  I look forward to contributing to a great company. I would love to talk more about it, please contact me at kiamcoo@gmail.com or call (909)-767-5806. You can also check out my portfolio at https://www.kiamco.dev/. Thank you for taking the time to read this, I hope to hear from you soon."
    
    doc = Document('cover_letter.docx')
    doc.add_paragraph(hi)
    doc.add_paragraph(first_paragraph)
    doc.add_paragraph(second_paragraph)
    doc.save('cover_letter.docx')

@click.command()
@click.option("-c","--company", help="company name" )
@click.option("-j","--job", help="job title")
@click.option("-f","--filetype", default="txt",help="ie: txt, docx, doc")
def cli(company, job, filetype):
    cover_template = f""" 
Hi,

    My name is Kim Kiamco,  I am excited for the opportunity to apply for the {job} position. I will be an asset to the team because of my commitment to building quality web applications. I am a part-time student at coding bootcamp (Lambda school) taking the Full-Stack Web Development track and will be finishing the program soon. Currently, I work as a part time Operations engineer at Menlo Security, a cyber security company. I have been looking for a full time position and thought {company} would be a great company to work for. My skills from Lambda school, working experience from Menlo Security and my adaptability in any environment will allow me to excel at {company}. My main tech stack when working on web applications is Javascript, React, Node.js, and PostgreSQL but I am willing and excited to learn new languages if needed. I am familiar with DevOps tools and practices like Docker, AWS and can write python scripts that can help with deployment automation.
 
    I have attached my resume and I believe that my experience and education would make me an asset at {company},  I look forward to contributing to a great company. I would love to talk more about it, please contact me at kiamcoo@gmail.com or call (909)-767-5806. You can also check out my portfolio at https://www.kiamco.dev/. Thank you for taking the time to read this, I hope to hear from you soon.   
 
Thanks, 
Kim Kiamco

    """
    # check if file type is docx 
    if filetype == 'docx':
        if_docx(job, company)

    # create file 
    file = open(f'cover_letter_for_{company}.{filetype}','w')
    file.write(cover_template)

    print("template created")
    print(cover_template)

if __name__ == "__main__":
    cli()